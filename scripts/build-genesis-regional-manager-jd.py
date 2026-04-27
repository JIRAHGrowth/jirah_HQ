"""
Build the Regional Manager – Metro Vancouver job description for Genesis Building Controls.

Output: Word .docx styled to approximate the existing Rascal HR-produced Genesis JDs
(navy headers, bullet lists, Page X of Y footer). No logo embedded — Joshua can either
paste into Rascal HR or Trent can drop the logo in once approved.

Structure:
  1. About the Company (Genesis boilerplate, verbatim)
  2. Core Values (Genesis boilerplate, verbatim)
  3. Role Summary
  4. Reporting & Relationships
  5. Key Responsibilities (grouped)
  6. Decision Rights
  7. Success Metrics
  8. Qualifications & Experience
  9. Work Environment & Travel
  10. Growth Pathway
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

# Genesis navy from the existing Rascal JDs (approximate match)
NAVY = RGBColor(0x1F, 0x3A, 0x6D)

LOGO_PATH = Path(__file__).parent / "assets" / "genesis-logo.png"

OUTPUT_DIR = Path(
    r"C:\Users\joshu\OneDrive - jirahgrowth.consulting\JIRAH Growth Partners - Shared"
    r"\01 - Clients\Active\Genesis Systems\07 - Deliverables"
)
OUTPUT_PATH = OUTPUT_DIR / "Regional-Manager-Metro-Vancouver.docx"


def add_page_number_footer(section):
    """Footer: right-aligned 'Page X of Y'."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)

    def field(instr_text):
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr_text)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")
        rPr.append(sz)
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = "1"
        run.append(t)
        fld.append(run)
        return fld

    run = p.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p._p.append(field("PAGE"))
    run2 = p.add_run(" of ")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p._p.append(field("NUMPAGES"))


def add_header_logo(section, logo_path: Path):
    """Place the Genesis logo top-left in the page header.

    Matches the existing Rascal HR-produced JDs (logo in header, appears on every
    page). Sized small to keep branding light.
    """
    if not logo_path.exists():
        return
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(logo_path), width=Inches(0.85))


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"


def add_subheading(doc, text):
    """Italic sub-header used inside KEY RESPONSIBILITIES groupings."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    run.font.name = "Calibri"


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"


def build():
    doc = Document()

    # Margins — matches Rascal look. Top margin slightly larger to give the header
    # logo room without crowding the title.
    for section in doc.sections:
        section.top_margin = Inches(1.1)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.35)
        add_header_logo(section, LOGO_PATH)
        add_page_number_footer(section)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ---------- TITLE ----------
    add_title(doc, "Regional Manager – Metro Vancouver")

    # ---------- ABOUT ----------
    add_heading(doc, "About the Company")
    add_body(
        doc,
        "Genesis Controls was started in 2013 as part of a large electrical contracting firm. "
        "As the controls division expanded, Genesis branched off into its own entity in 2017. "
        "With over 20 years of experience in the controls trade, Genesis is quickly becoming "
        "one of the largest controls companies in BC. With our expansion into Metro Vancouver "
        "and Washington state, we continue to grow and add to our portfolio. Our relationships "
        "with contractors and the community have been and will continue to be the key to our "
        "success. Genesis Controls is a progressive, growing organization committed to exceeding "
        "our clients' expectations by providing one-stop exceptional services for building "
        "integration and controls. With experience in hi-rise, commercial, institutional, "
        "residential, industrial, and healthcare sectors, our proven track record has provided "
        "professional service and maximum value to our clients.",
    )

    # ---------- CORE VALUES ----------
    add_heading(doc, "Core Values")
    add_body(
        doc,
        "We will provide outstanding customer service through personalized, professional "
        "excellence, guaranteed workmanship, innovative technology, competitive pricing and "
        "efficient scheduling while maintaining the highest level of safety standards and "
        "environmental protection at all times. We will always strive to provide maximum value "
        "and establish long term relationships with our clients, suppliers and staff through "
        "integrity, dedication and ethical business practices.",
    )

    # ---------- ROLE SUMMARY ----------
    add_heading(doc, "Role Summary")
    add_body(
        doc,
        "The Regional Manager – Metro Vancouver is the senior operational leader for the "
        "Metro Vancouver office and carries single-point accountability for the region's "
        "people, projects, and performance. Reporting directly to the Owner, this role is the "
        "bridge between Kelowna headquarters and the Metro Vancouver team — translating "
        "company direction into regional execution, and bringing regional realities back into "
        "company-wide decisions. It is the first in a family of regional leadership roles at "
        "Genesis and sets the standard for how regional offices are led as the company grows.",
    )

    # ---------- REPORTING & RELATIONSHIPS ----------
    add_heading(doc, "Reporting & Relationships")
    add_bullets(
        doc,
        [
            "Reports to: Owner / Chief Executive Officer.",
            "Direct reports: Project Managers, Controls Technologists, Controls Electricians "
            "(Journeyperson and Apprentice), and the regional Service Technician assigned to "
            "Metro Vancouver.",
            "Works alongside the Office Manager (headquarters-based — finance, accounts "
            "payable and receivable, payroll, group benefits, and HR administration for all "
            "regions). This function serves Metro Vancouver but does not report through the "
            "Regional Manager.",
            "Works alongside the Lead Estimator (headquarters-based — bid strategy, "
            "subcontractor relationships, and estimating workload for all regions). This "
            "function supports Metro Vancouver project pursuits but does not report through "
            "the Regional Manager.",
            "Coordinates closely with the Business Development Manager on regional pipeline, "
            "key account relationships, and bid follow-through.",
            "Partners with headquarters leadership on cross-region staffing, resource "
            "allocation, and company-wide initiatives.",
        ],
    )

    # ---------- KEY RESPONSIBILITIES ----------
    add_heading(doc, "Key Responsibilities")

    add_subheading(doc, "People & Culture")
    add_bullets(
        doc,
        [
            "Lead, coach, and develop the Metro Vancouver team across project management, "
            "technology, trades, and service functions.",
            "Own regional hiring, onboarding, and performance conversations in partnership "
            "with the Owner and the Office Manager.",
            "Build and sustain a strong regional culture consistent with Genesis's "
            "company-wide values, with particular attention to field-based and distributed "
            "team members.",
            "Serve as the primary leadership presence for regional staff and the day-to-day "
            "bridge to headquarters for questions, concerns, and escalations.",
        ],
    )

    add_subheading(doc, "Operations & Delivery")
    add_bullets(
        doc,
        [
            "Oversee execution of all projects and service work delivered from the Metro "
            "Vancouver office, ensuring quality, schedule, and budget commitments are met.",
            "Enforce and continually improve Genesis's documented SOPs — bid-to-kickoff, "
            "project handoff, service dispatch, and project close-out — at the regional level.",
            "Coordinate with the Lead Estimator and Project Managers on project kickoff, "
            "scope clarity, and clean handoffs into field execution.",
            "Monitor jobsite safety, compliance with provincial regulations, and safety "
            "culture across the region.",
            "The Regional Manager does not carry a dedicated service route or a routine "
            "hands-on service workload. Field response is limited to escalations and "
            "exception cases where leadership presence is required.",
        ],
    )

    add_subheading(doc, "Client & Market")
    add_bullets(
        doc,
        [
            "Maintain relationships with Metro Vancouver general contractors, consulting "
            "engineers, building owners, and service clients.",
            "Represent Genesis at regional industry events (MCA, VRCA, and similar) and act "
            "as the senior face of the company in the region.",
            "Work with the Business Development Manager on regional pipeline development, "
            "key account strategy, and bid-award follow-through.",
            "Identify regional market trends, competitive activity, and growth opportunities "
            "and surface them to the Owner and leadership team.",
        ],
    )

    add_subheading(doc, "Financial & Reporting")
    add_bullets(
        doc,
        [
            "Own regional operational performance against agreed KPIs, including regional "
            "gross margin, project profitability, service responsiveness, and staff "
            "utilization.",
            "Monitor regional job costing in partnership with the Office Manager and Project "
            "Managers; flag margin erosion, scope creep, and cash flow risks early.",
            "Prepare a monthly regional performance summary for the Owner and contribute to "
            "quarterly leadership reviews.",
            "Provide input to the annual regional budgeting and workforce planning process.",
        ],
    )

    # ---------- DECISION RIGHTS ----------
    add_heading(doc, "Decision Rights")
    add_body(
        doc,
        "The Regional Manager carries authority to make the following decisions "
        "independently, without escalation:",
    )
    add_bullets(
        doc,
        [
            "Day-to-day scheduling, project assignment, and field resource allocation within "
            "the region.",
            "Hiring decisions for trade and technologist roles within approved headcount; "
            "termination and disciplinary decisions for regional staff in consultation with "
            "the Owner.",
            "Operational purchases and commitments up to a regional spending limit to be "
            "established with the Owner.",
            "Regional client issue resolution and service recovery.",
            "Subcontractor engagement on regional projects, within pricing and scope "
            "parameters set by the Lead Estimator.",
        ],
    )

    add_body(doc, "The following decisions are owned by the Owner and must be escalated:")
    add_bullets(
        doc,
        [
            "Headcount changes beyond approved plan, executive hires, or senior role "
            "compensation.",
            "Project commitments exceeding the regional contract-authority threshold.",
            "Changes to company pricing, service packaging, or formal commercial policy.",
            "Major capital investments, vehicle purchases, or office lease decisions.",
            "Any matter carrying material financial, legal, or reputational exposure.",
        ],
    )

    add_body(doc, "The following are owned jointly with peers:")
    add_bullets(
        doc,
        [
            "Bid strategy and go/no-go decisions — with the Lead Estimator.",
            "Payroll, benefits, and HR-policy application — with the Office Manager.",
            "Cross-region staff transfers or resource borrowing — with counterpart regional "
            "leaders as they are established.",
        ],
    )

    # ---------- SUCCESS METRICS ----------
    add_heading(doc, "Success Metrics")
    add_body(
        doc,
        "Regional performance is reviewed monthly against the following measurables:",
    )
    add_bullets(
        doc,
        [
            "Regional gross margin against target.",
            "Project on-time and on-budget completion rate.",
            "Service call response time (from request to dispatch) and first-visit resolution "
            "rate.",
            "Regional employee retention and voluntary turnover rate.",
            "Field utilization and overtime trend.",
            "Safety incident rate and near-miss reporting participation.",
            "Regional contribution to the company's growth toward the $10M revenue milestone.",
        ],
    )

    # ---------- QUALIFICATIONS ----------
    add_heading(doc, "Qualifications & Experience")

    add_subheading(doc, "Required")
    add_bullets(
        doc,
        [
            "Minimum 8 years of operational experience in building controls, BMS, mechanical "
            "contracting, or an adjacent construction trade, with at least 3 years in a "
            "formal supervisory or operations role.",
            "Demonstrated ability to lead cross-disciplinary teams including project "
            "managers, technologists, trades, and service personnel.",
            "Strong working knowledge of HVAC systems, building automation, and BC "
            "electrical and building codes and safety regulations.",
            "Ability to read and interpret construction documents, project schedules, and "
            "financial statements at a P&L line-item level.",
            "Excellent written and verbal communication, with the ability to translate "
            "between field, engineering, and executive audiences.",
            "Valid BC driver's license and access to reliable transportation.",
        ],
    )

    add_subheading(doc, "Preferred")
    add_bullets(
        doc,
        [
            "Journeyperson Electrician certification or equivalent controls-trade credential.",
            "Post-secondary education in construction management, engineering, business, or "
            "a related field.",
            "Prior experience managing a geographically distinct office or region.",
            "Familiarity with project management and job-costing platforms (SiteMax, "
            "Procore, Sage, or equivalent).",
            "Track record of culture-building in a field-based, distributed team.",
        ],
    )

    add_subheading(doc, "Personal Attributes")
    add_bullets(
        doc,
        [
            "Calm and decisive under operational pressure.",
            "Owner-mindset — acts as a steward of the business, not as a task manager.",
            "Willing and able to shift focus from hands-on trade work into full-time "
            "management.",
            "Genuinely committed to developing the people on the regional team.",
            "Aligned with Genesis's core values and able to carry them into the region.",
        ],
    )

    # ---------- WORK ENVIRONMENT ----------
    add_heading(doc, "Work Environment & Travel")
    add_bullets(
        doc,
        [
            "Based at the Metro Vancouver office.",
            "Regular jobsite presence across Metro Vancouver and the Fraser Valley.",
            "Quarterly travel to Kelowna headquarters for leadership reviews and team events "
            "(approximately one week per quarter).",
            "Occasional travel to other Genesis regional offices as the company expands.",
        ],
    )

    # ---------- GROWTH PATHWAY ----------
    add_heading(doc, "Growth Pathway")
    add_body(
        doc,
        "The Regional Manager role sits on Genesis's senior operational leadership track. "
        "It is the prototype for regional leadership as the company expands into Alberta "
        "and beyond. Strong performance in this role positions the incumbent for broader "
        "operational, multi-region, or senior leadership responsibility as Genesis grows.",
    )

    # ---------- SAVE ----------
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
