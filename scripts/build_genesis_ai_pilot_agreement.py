"""Generate the Genesis AI Pilot Agreement draft + Counsel Cover Memo.

Outputs two .docx files to the Genesis Admin/Contract & DocuSign folder:
  - Genesis_JIRAH AI Pilot Agreement - DRAFT FOR COUNSEL.docx
  - Genesis_JIRAH AI Pilot - Cover Memo to Counsel.docx

Reflects the decisions captured in the April 23, 2026 pre-counsel discussion:
  - Standalone agreement (not MSA+SOW for now; flagged in cover memo)
  - 3-month pilot, $9,095/mo, Day-2 Plan A/B gate + Month 3 decision gate
  - Client owns final Tool; JIRAH retains Reusable Components
  - Client sets up own third-party API accounts (own keys at handoff)
  - Zero-hallucination evaluation harness = process commitment, not outcome warranty
  - Two-tier change management (minor log / material amendment)
  - Continued Operator Services = separate SOW option
  - Liability cap: floor $50k / 12 months' fees / ceiling $150k (placeholders pending broker)
  - In-contract PIPEDA language (no separate DPA schedule)
  - Subprocessor list as Schedule B with 15-day objection right
  - Publicity anonymization with approval
"""

from __future__ import annotations

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).parent))
from _paths import ACTIVE_CLIENTS, assert_no_legacy_segment

BASE_DIR = str(
    ACTIVE_CLIENTS
    / "Genesis Systems"
    / "03 - Admin"
    / "Contract & DocuSign"
)
assert_no_legacy_segment(BASE_DIR)
AGREEMENT_PATH = os.path.join(BASE_DIR, "Genesis_JIRAH AI Pilot Agreement - DRAFT FOR COUNSEL.docx")
MEMO_PATH = os.path.join(BASE_DIR, "Genesis_JIRAH AI Pilot - Cover Memo to Counsel.docx")


def new_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    return doc


def _h1(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def _h2(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def _h3(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 3")


def _p(doc: Document, text: str = "", bold: bool = False, italic: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _draft_banner(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run("DRAFT — FOR COUNSEL REVIEW · NOT FOR EXECUTION")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


# ---------------------------------------------------------------------------
# Agreement
# ---------------------------------------------------------------------------

def build_agreement() -> None:
    doc = new_document()

    _draft_banner(doc)
    _p(doc, "Prepared: April 23, 2026")
    _p(doc, (
        "Bracketed items [like this] indicate placeholders requiring confirmation by the "
        "parties, by counsel, or pending the broker review of JIRAH's insurance coverage. A "
        "separate Counsel Cover Memo accompanies this draft summarising open items flagged "
        "for counsel's specific attention."
    ), italic=True)
    _p(doc, "")

    _h1(doc, "AI Pilot Development Agreement")

    _p(doc, "Effective Date: April 21, 2026")
    _p(doc, "")
    _p(doc, "JIRAH Growth Consulting Ltd.")
    _p(doc, "3610 Weston Rd, Kelowna, BC V1X 6K2 (“Consultant”)")
    _p(doc, "")
    _p(doc, "And")
    _p(doc, "")
    _p(doc, "Genesis Building Controls Ltd.")
    _p(doc, "[Registered address — to be confirmed] (“Client”)")

    # -------- Recitals --------
    _h2(doc, "Recitals")
    _p(doc, (
        "A. The parties have entered into a Strategic Execution Advisory Agreement dated "
        "April 21, 2026 (the “Strategic Agreement”) governing Consultant's strategic advisory "
        "services to Client."
    ))
    _p(doc, (
        "B. The parties now wish to enter into this separate, parallel agreement for an AI "
        "Pilot engagement, under which Consultant will design and build AI-enabled tooling "
        "for Client's estimating and project-management workflows."
    ))
    _p(doc, (
        "C. This Agreement operates independently of the Strategic Agreement as further set "
        "out in Section 25."
    ))

    # -------- 1. Definitions --------
    _h2(doc, "1. Definitions")
    _p(doc, "In this Agreement:")
    _bullet(doc, (
        "“Client Data” means all data, information, and materials provided by Client to "
        "Consultant for the performance of the Services, including without limitation Client's "
        "estimating data, project-management data, pricing data, customer data, and employee "
        "personal information."
    ))
    _bullet(doc, (
        "“Evaluation Harness” means the testing process and acceptance criteria jointly "
        "defined by the parties under Section 5."
    ))
    _bullet(doc, (
        "“Pilot Term” means the three (3) month period commencing on the Effective Date and "
        "ending July 21, 2026, unless earlier terminated or extended in accordance with this "
        "Agreement."
    ))
    _bullet(doc, (
        "“Reusable Components” means Consultant's proprietary prompt frameworks, orchestration "
        "templates, evaluation methodologies, and abstracted patterns that do not contain "
        "Client-specific data, terminology, configurations, or identifiers. Examples include, "
        "without limitation, multi-stage validation prompt structures, evaluation-harness "
        "methodology, orchestration templates, abstracted workflow patterns, and anonymised "
        "lessons learned."
    ))
    _bullet(doc, (
        "“Subprocessors” means the third-party service providers listed in Schedule B, as "
        "amended from time to time in accordance with Section 11."
    ))
    _bullet(doc, (
        "“Tool” means the integrated AI-enabled software artifacts custom-built by Consultant "
        "for Client under this Agreement, including Client-specific configurations, prompts, "
        "workflows, integrations, and deployment. For clarity, the Tool does not include the "
        "Reusable Components incorporated into it, which remain Consultant's property subject "
        "to Section 10."
    ))
    _bullet(doc, (
        "“Use Cases” means (i) Client's estimating workflow and (ii) Client's project-"
        "management workflow, as further described in Schedule A, and no others."
    ))

    # -------- 2. Engagement --------
    _h2(doc, "2. Engagement")
    _p(doc, (
        "Client retains Consultant to design, build, and deliver the Tool for Client's use in "
        "the Use Cases, on a pilot basis, in accordance with the terms of this Agreement. The "
        "parties acknowledge that this is a pilot engagement with defined decision gates "
        "(Section 4) and that continued development of the Tool beyond the Pilot Term is "
        "contingent on the outcome of those gates and the parties' mutual agreement."
    ))

    # -------- 3. Scope of Services --------
    _h2(doc, "3. Scope of Services")
    _p(doc, "During the Pilot Term, Consultant will:")
    _bullet(doc, "Design and build the Tool for the Use Cases identified in Schedule A;")
    _bullet(doc, (
        "Conduct weekly working sessions with Client's nominated champion and Client's "
        "leadership team (virtual or in-person, as mutually agreed);"
    ))
    _bullet(doc, (
        "Deliver a demonstrable artifact at the end of each week of the Pilot Term (each a "
        "“Weekly Demo”);"
    ))
    _bullet(doc, "Define and operate the Evaluation Harness in accordance with Section 5;")
    _bullet(doc, (
        "Integrate the Tool with Client's systems as required and mutually agreed in writing;"
    ))
    _bullet(doc, (
        "Transition the Tool's operational configuration to Client's Subprocessor accounts at "
        "handoff in accordance with Sections 11 and 19;"
    ))
    _bullet(doc, "Train Client's internal champion on Tool operation per Section 19.")

    _p(doc, (
        "Except as expressly agreed in writing under a separate statement of work or amendment, "
        "Consultant will NOT:"
    ))
    _bullet(doc, (
        "Provide production operations, service-level commitments, uptime warranties, or 24/7 "
        "support;"
    ))
    _bullet(doc, "Develop use cases beyond the Use Cases specified in Schedule A;")
    _bullet(doc, "Conduct ongoing model retraining or maintenance after the Pilot Term;")
    _bullet(doc, "Train personnel beyond Client's nominated internal champion;")
    _bullet(doc, "Develop or operate a multi-tenant software-as-a-service product;")
    _bullet(doc, (
        "Provide any of the services contemplated under the Strategic Agreement, which remain "
        "governed solely by that agreement."
    ))

    # -------- 4. Decision Gates --------
    _h2(doc, "4. Decision Gates")
    _p(doc, "The Pilot is structured around two formal decision gates:")

    _h3(doc, "(a) Day-2 Plan A / Plan B Gate")
    _p(doc, (
        "Within the first two weeks of the Pilot Term, Consultant will assess the feasibility "
        "of the primary technical approach (“Plan A”) against the feasibility criteria "
        "documented at the outset of the Pilot. If Plan A is not feasible, Consultant will "
        "pivot to the fallback approach (“Plan B”) or recommend termination of the Pilot. The "
        "outcome of this gate will be documented in writing and delivered to Client within "
        "three business days of the assessment."
    ))

    _h3(doc, "(b) Month 3 Decision Gate")
    _p(doc, (
        "At the end of the Pilot Term, the parties will jointly review the Tool's performance "
        "against the Evaluation Harness acceptance criteria and elect one of the following "
        "outcomes:"
    ))
    _bullet(doc, (
        "(i) Handoff — the Tool is handed off to Client in accordance with Section 19 and "
        "this Agreement terminates in accordance with Section 17;"
    ))
    _bullet(doc, (
        "(ii) Extension — the parties enter a separate written agreement for continued "
        "development, operator services, or further pilot phases on mutually agreed terms; or"
    ))
    _bullet(doc, (
        "(iii) Termination Without Handoff — the Pilot ends without handoff; Consultant "
        "delivers available artifacts in their then-current state in accordance with "
        "Section 17(e); each party performs its post-termination obligations."
    ))
    _p(doc, (
        "Neither party is obligated to elect any particular outcome. Absent mutual election "
        "within fifteen (15) days of the end of the Pilot Term, the default outcome is (i) "
        "Handoff, unless Client elects (iii) Termination Without Handoff by written notice."
    ))

    # -------- 5. Evaluation Harness --------
    _h2(doc, "5. Evaluation Harness")
    _p(doc, (
        "Within the first two weeks of the Pilot Term, the parties will jointly document the "
        "Evaluation Harness, which shall include:"
    ))
    _bullet(doc, "Test cases representative of Client's operational scenarios;")
    _bullet(doc, "Acceptance criteria defining what constitutes acceptable Tool performance;")
    _bullet(doc, "Testing cadence, documentation protocol, and remediation process.")
    _p(doc, (
        "Consultant commits to operating the Evaluation Harness against each significant Tool "
        "release prior to Client deployment, and to documenting results and remediation "
        "actions."
    ))
    _p(doc, (
        "The Evaluation Harness is a process commitment, not a warranty of output accuracy. "
        "Consultant does not warrant that the Tool will produce accurate or error-free outputs "
        "in production, and the Tool's outputs shall not be relied upon without human review "
        "as provided in Section 14."
    ))

    # -------- 6. Client Obligations --------
    _h2(doc, "6. Client Obligations")
    _p(doc, "Client will:")
    _bullet(doc, (
        "Nominate an internal champion within fourteen (14) days of the Effective Date and "
        "ensure that person's availability for weekly working sessions and handoff training;"
    ))
    _bullet(doc, (
        "Provide Consultant with reasonable access to Client Data, systems, and stakeholders "
        "necessary for performance of the Services;"
    ))
    _bullet(doc, (
        "Establish and maintain accounts with the Subprocessors listed in Schedule B as "
        "required for the Tool's operation at and after handoff, and bear all usage fees "
        "associated with such accounts from handoff forward;"
    ))
    _bullet(doc, (
        "Review and approve Weekly Demos and key milestone deliverables in a timely manner;"
    ))
    _bullet(doc, (
        "Respond to Consultant's data or clarification requests within five (5) business days, "
        "or as otherwise mutually agreed in writing;"
    ))
    _bullet(doc, (
        "Warrant that it has provided all notices and obtained all consents required under "
        "applicable privacy law for the Client Data provided to Consultant, including without "
        "limitation any consents required for processing of Client's employees' personal "
        "information."
    ))
    _p(doc, (
        "Client acknowledges that Consultant's ability to meet timelines and milestones is "
        "dependent on Client's timely performance of these obligations. Consultant will not "
        "be in breach of this Agreement on account of delay directly attributable to Client's "
        "non-performance of its obligations."
    ))

    # -------- 7. Fees and Payment --------
    _h2(doc, "7. Fees and Payment")
    _bullet(doc, "Monthly fee: $9,095 CAD plus applicable GST.")
    _bullet(doc, (
        "Payable monthly in advance upon receipt of invoice. Total Pilot Term fees: $27,285 "
        "CAD plus GST, exclusive of any third-party usage overage under this Section."
    ))
    _bullet(doc, (
        "Third-Party Usage Ceiling. The monthly fee includes up to [$800] CAD per month in "
        "third-party AI service usage paid through Consultant's development accounts during "
        "the Pilot Term. Usage exceeding this ceiling will be invoiced to Client at cost with "
        "supporting documentation."
    ))
    _bullet(doc, (
        "Late payments exceeding thirty (30) days may result in suspension of the Services "
        "and/or interest charged at 2% per month (24% per annum)."
    ))
    _bullet(doc, "No onboarding or setup fees apply.")
    _bullet(doc, (
        "Fees are fixed for the Pilot Term. Any extension, continuation, or expansion of the "
        "engagement requires a separate written agreement with its own fee terms."
    ))

    # -------- 8. Term --------
    _h2(doc, "8. Term")
    _p(doc, (
        "This Agreement commences on April 21, 2026 (the “Effective Date”) and continues for "
        "the Pilot Term of three (3) months, ending July 21, 2026, unless earlier terminated "
        "in accordance with Section 17 or continued by separate written agreement pursuant to "
        "Section 4(b)."
    ))

    # -------- 9. Meeting Expectations --------
    _h2(doc, "9. Meeting Expectations")
    _p(doc, (
        "Consultant and Client will conduct weekly working sessions (virtual or in-person, as "
        "mutually agreed session-by-session), a mid-Pilot checkpoint review, and a Month 3 "
        "Decision Gate review. Consultant may include additional members of the JIRAH team "
        "(e.g., co-founder or subject-matter contributors) in relevant sessions at no "
        "additional cost to Client. Consultant will remain available for reasonable "
        "communication via email and/or phone between sessions."
    ))

    # -------- 10. IP --------
    _h2(doc, "10. Intellectual Property")

    _p(doc, "(a) Reusable Components.", bold=True)
    _p(doc, (
        "All Reusable Components are and remain the sole and exclusive property of Consultant. "
        "Upon payment of all fees due under this Agreement, Client receives a perpetual, "
        "non-exclusive, non-transferable, royalty-free licence to use Reusable Components "
        "solely as incorporated into the Tool and solely for Client's internal business "
        "purposes. Consultant retains the right to use, modify, and further develop Reusable "
        "Components for its own purposes and for its other clients."
    ))

    _p(doc, "(b) The Tool.", bold=True)
    _p(doc, (
        "Upon the later of (i) payment in full of all Pilot Term fees and (ii) handoff "
        "pursuant to Section 19, Client owns the Tool, including all Client-specific "
        "configurations, prompts, workflows, integrations, and deployment artifacts developed "
        "for Client. Consultant hereby assigns to Client all right, title, and interest in "
        "and to such Client-specific components of the Tool, subject only to Consultant's "
        "retained rights in the Reusable Components under Section 10(a)."
    ))

    _p(doc, "(c) AI-Generated Outputs.", bold=True)
    _p(doc, (
        "The parties acknowledge that copyright protection for AI-generated outputs is "
        "unsettled under Canadian law. Neither party asserts copyright ownership in outputs "
        "generated by the Tool. Client may use such outputs without restriction as between "
        "the parties, subject to the warranties, disclaimers, and human-in-the-loop "
        "requirements in Section 14."
    ))

    _p(doc, "(d) Client Materials.", bold=True)
    _p(doc, (
        "All Client Data, Client-provided materials, and Client's pre-existing intellectual "
        "property remain the sole property of Client. Consultant receives a limited, "
        "non-exclusive, royalty-free licence to use Client Data solely to perform the "
        "Services during the Pilot Term and for no other purpose."
    ))

    _p(doc, "(e) Third-Party Intellectual Property.", bold=True)
    _p(doc, (
        "Each party warrants that the materials and data it provides under this Agreement do "
        "not, to the best of its knowledge, infringe the intellectual property rights of any "
        "third party."
    ))

    _p(doc, (
        "Additional detail on the allocation of ownership and licence rights, including "
        "worked examples, is set out in Schedule A (Intellectual Property Matrix)."
    ))

    # -------- 11. Subprocessors --------
    _h2(doc, "11. Third-Party Service Providers and Subprocessors")

    _p(doc, "(a) Initial Subprocessor List.", bold=True)
    _p(doc, (
        "The third-party service providers initially approved for use in performing the "
        "Services are listed in Schedule B, together with their respective purposes and data-"
        "residency regions."
    ))

    _p(doc, "(b) Addition of Subprocessors.", bold=True)
    _p(doc, (
        "Consultant will provide Client with at least fifteen (15) days' written notice prior "
        "to engaging a new Subprocessor. Client may object to such addition on reasonable "
        "grounds within that period. Failure to object within the period shall constitute "
        "Client's consent. If Client objects and the parties cannot agree on a resolution "
        "within a further fifteen (15) days, Client's sole remedy is to terminate this "
        "Agreement on written notice."
    ))

    _p(doc, "(c) Client Accounts at and after Handoff.", bold=True)
    _p(doc, (
        "Client will establish and maintain accounts with each Subprocessor required for the "
        "Tool's production operation. Client bears all usage fees associated with such "
        "accounts from handoff forward."
    ))

    _p(doc, "(d) Transition at Handoff.", bold=True)
    _p(doc, (
        "At handoff, Consultant will transition the Tool's operational configuration to "
        "Client's Subprocessor accounts and will revoke its operational access, except as "
        "may be required under a separately agreed engagement under Section 20."
    ))

    _p(doc, "(e) Third-Party Terms.", bold=True)
    _p(doc, (
        "Client's use of Subprocessors is subject to those Subprocessors' then-current terms "
        "of service. Changes to Subprocessor terms, availability, pricing, or service levels "
        "imposed by those Subprocessors are outside Consultant's control, and Consultant's "
        "liability for resulting changes or disruptions is limited as provided in Section 16."
    ))

    # -------- 12. Data Protection --------
    _h2(doc, "12. Data Protection and Privacy")

    _p(doc, "(a) Roles.", bold=True)
    _p(doc, (
        "For the purposes of applicable privacy law, including the Personal Information "
        "Protection and Electronic Documents Act (Canada) and the Personal Information "
        "Protection Act (British Columbia), Client is the controller of Client Data and "
        "Consultant processes Client Data on Client's behalf as service provider."
    ))

    _p(doc, "(b) Purpose Limitation.", bold=True)
    _p(doc, (
        "Consultant will process Client Data only to perform the Services under this "
        "Agreement and for no other purpose. Without limiting the foregoing, Consultant will "
        "not use Client Data: (i) to train general-purpose AI models; (ii) in case studies or "
        "marketing materials except as expressly permitted and anonymised under Section 22; "
        "or (iii) for benchmarking across Consultant's other engagements."
    ))

    _p(doc, "(c) Client Consents.", bold=True)
    _p(doc, (
        "Client warrants that it has provided all required notices and obtained all required "
        "consents under applicable privacy law for the Client Data provided to Consultant, "
        "including without limitation any notices or consents required for processing of "
        "Client's employees' personal information."
    ))

    _p(doc, "(d) Data Residency.", bold=True)
    _p(doc, (
        "Client Data will be processed only in Canada or the United States. Consultant will "
        "not process Client Data in any other jurisdiction without Client's prior written "
        "consent."
    ))

    _p(doc, "(e) Security Safeguards.", bold=True)
    _p(doc, (
        "Consultant will implement and maintain administrative, technical, and physical "
        "safeguards appropriate to the sensitivity of Client Data, consistent with industry "
        "standards for comparable services."
    ))

    _p(doc, "(f) Personnel Access.", bold=True)
    _p(doc, (
        "Access to Client Data will be limited to Consultant personnel with a need to know "
        "for performance of the Services, and such personnel are bound by confidentiality "
        "obligations at least as restrictive as those in Section 13."
    ))

    _p(doc, "(g) Breach Notification.", bold=True)
    _p(doc, (
        "Consultant will notify Client without undue delay, and in any event within seventy-"
        "two (72) hours of becoming aware, of any unauthorised access to, disclosure of, or "
        "loss of Client Data, and will cooperate reasonably with Client's investigation and "
        "response."
    ))

    _p(doc, "(h) Retention and Deletion.", bold=True)
    _p(doc, (
        "Upon termination or expiry of this Agreement, Consultant will, at Client's written "
        "election, return Client Data to Client or securely delete Client Data from "
        "Consultant's systems, except to the extent retention is required by applicable law. "
        "Consultant will provide written certification of deletion upon Client's request."
    ))

    _p(doc, "(i) Subprocessor Diligence.", bold=True)
    _p(doc, (
        "Consultant will engage Subprocessors only under terms substantially consistent with "
        "this Section 12."
    ))

    # -------- 13. Confidentiality --------
    _h2(doc, "13. Confidentiality")
    _p(doc, (
        "Each party will maintain the confidentiality of all proprietary, personal, or "
        "business information of the other party disclosed during the course of this "
        "Agreement, whether identified as confidential or reasonably understood to be "
        "confidential given the nature of the information or the circumstances of "
        "disclosure. Each party will take reasonable precautions to protect such information "
        "and will not disclose it to any third party without prior written consent, except as "
        "required by law or legal process. This Section survives termination."
    ))

    # -------- 14. AI Warranties & Disclaimers --------
    _h2(doc, "14. AI Warranties and Disclaimers")

    _p(doc, "(a) Services Standard.", bold=True)
    _p(doc, (
        "Consultant warrants that it will perform the Services with the professional skill "
        "and care consistent with industry standards for AI-pilot development work."
    ))

    _p(doc, "(b) No Warranty of AI Accuracy.", bold=True)
    _p(doc, (
        "Client acknowledges that AI technologies may produce inaccurate, incomplete, "
        "inappropriate, or unexpected outputs. Consultant does not warrant that the Tool will "
        "produce accurate or reliable outputs in any or all circumstances. The Evaluation "
        "Harness (Section 5) is a process commitment and not a warranty of output."
    ))

    _p(doc, "(c) Human-in-the-Loop Requirement.", bold=True)
    _p(doc, (
        "Client will maintain human review of Tool outputs prior to any use with material "
        "business consequence, including without limitation the submission of bids, "
        "communications with customers, or decisions regarding employees. Client retains "
        "final authority over all such uses and over any decision to act on Tool outputs."
    ))

    _p(doc, "(d) No Warranty of Fitness.", bold=True)
    _p(doc, (
        "Except as expressly set out in this Agreement, Consultant makes no warranties, "
        "express or implied, including without limitation any implied warranty of "
        "merchantability, fitness for a particular purpose, or non-infringement."
    ))

    _p(doc, "(e) Reliance at Client's Risk.", bold=True)
    _p(doc, (
        "Client agrees that it relies on the Tool's outputs at its own discretion and risk. "
        "Consultant is not liable for business decisions made by Client or third parties based "
        "on Tool outputs."
    ))

    # -------- 15. Insurance --------
    _h2(doc, "15. Insurance")
    _p(doc, (
        "Consultant will maintain, at its expense, during the Pilot Term and for a period of "
        "at least twelve (12) months thereafter: (i) Errors & Omissions (professional "
        "liability) insurance; and (ii) cyber liability insurance, in each case in amounts "
        "consistent with industry practice for services of comparable scope. Consultant will "
        "provide certificates of insurance upon Client's written request."
    ))

    # -------- 16. Liability & Indemnification --------
    _h2(doc, "16. Liability and Indemnification")

    _p(doc, "(a) Exclusion of Damages.", bold=True)
    _p(doc, (
        "In no event will either party be liable for any indirect, incidental, consequential, "
        "special, exemplary, or punitive damages, including without limitation loss of "
        "revenue, profit, data, goodwill, or business opportunity, arising out of or in "
        "connection with this Agreement or the Services, whether in contract, tort, or "
        "otherwise, even if advised of the possibility of such damages."
    ))

    _p(doc, "(b) Aggregate Liability Cap.", bold=True)
    _p(doc, (
        "Except for the Excluded Claims defined in Section 16(c), each party's aggregate "
        "liability arising out of or in connection with this Agreement will be limited to the "
        "greater of (i) the fees paid by Client to Consultant under this Agreement in the "
        "twelve (12) months preceding the event giving rise to the claim, or (ii) [$50,000] "
        "CAD. In no event will either party's aggregate liability under this Agreement exceed "
        "[$150,000] CAD."
    ))
    _p(doc, (
        "[DRAFTING NOTE FOR COUNSEL: the dollar floor and ceiling are placeholders pending "
        "confirmation of Consultant's Errors & Omissions and cyber liability coverage limits. "
        "The ceiling should be set at or below available coverage so uninsured exposure is "
        "zero.]"
    ), italic=True)

    _p(doc, "(c) Excluded Claims.", bold=True)
    _p(doc, (
        "The limitations in Section 16(a) and (b) do not apply to: (i) Client's payment "
        "obligations under Section 7; (ii) breach of the confidentiality obligations in "
        "Section 13; (iii) indemnification obligations under Sections 16(d) and 16(e); or "
        "(iv) liability arising from a party's gross negligence or wilful misconduct."
    ))

    _p(doc, "(d) Consultant Indemnification.", bold=True)
    _p(doc, (
        "Consultant will indemnify, defend, and hold harmless Client from and against "
        "third-party claims to the extent arising from: (i) Consultant's breach of this "
        "Agreement; (ii) Consultant's gross negligence or wilful misconduct; or (iii) "
        "infringement by the Tool (excluding Client Materials, Client's modifications, and "
        "Subprocessor outputs) of a third party's intellectual property rights."
    ))

    _p(doc, "(e) Client Indemnification.", bold=True)
    _p(doc, (
        "Client will indemnify, defend, and hold harmless Consultant from and against "
        "third-party claims to the extent arising from: (i) Client's breach of this "
        "Agreement; (ii) Client's failure to obtain required notices or consents for the "
        "Client Data provided under Section 6; (iii) Client's use of Tool outputs in "
        "violation of Section 14; or (iv) infringement by Client Materials, or by Client's "
        "use of the Tool in a manner not contemplated by this Agreement, of a third party's "
        "intellectual property rights."
    ))

    _p(doc, "(f) Subprocessor Pass-Through.", bold=True)
    _p(doc, (
        "Client acknowledges that the Services require use of third-party AI service "
        "providers and other Subprocessors, that breaches, failures, or changes by such "
        "providers may occur outside Consultant's direct control, and that Consultant's "
        "liability for losses arising from acts or omissions of Subprocessors is limited to "
        "the amounts actually recovered by Consultant from such Subprocessors, subject in all "
        "events to the caps in Section 16(b)."
    ))

    # -------- 17. Termination --------
    _h2(doc, "17. Termination")
    _p(doc, (
        "(a) Either party may terminate this Agreement by providing thirty (30) days' written "
        "notice to the other party."
    ))
    _p(doc, (
        "(b) Either party may terminate this Agreement immediately by written notice if the "
        "other party materially breaches this Agreement and fails to cure the breach within "
        "fifteen (15) days of written notice of the breach."
    ))
    _p(doc, (
        "(c) This Agreement terminates automatically at the end of the Pilot Term unless "
        "continued or extended by separate written agreement pursuant to Section 4(b)."
    ))
    _p(doc, (
        "(d) All fees and expenses incurred up to the effective date of termination remain "
        "payable."
    ))
    _p(doc, (
        "(e) Upon termination for any reason, Consultant will, upon Client's written request: "
        "(i) deliver the Tool to Client in its then-current state, together with then-current "
        "documentation; (ii) cooperate reasonably with transition to Client or Client's "
        "designee; and (iii) comply with Section 12(h) regarding Client Data."
    ))
    _p(doc, (
        "(f) Sections 10, 12, 13, 14, 16, 20, 21, 22, 23, 24, and 25 survive termination of "
        "this Agreement."
    ))

    # -------- 18. Change Management --------
    _h2(doc, "18. Change Management")

    _p(doc, "(a) Minor Changes.", bold=True)
    _p(doc, (
        "Adjustments within the Use Cases, UI refinements, additional test cases, and similar "
        "minor changes that do not materially increase scope, effort, or cost may be "
        "implemented on the basis of email confirmation from Client's champion and will be "
        "logged in a shared change log. Minor changes shall remain within reasonable effort "
        "commensurate with the Pilot Term and do not alter the fees payable under Section 7."
    ))

    _p(doc, "(b) Material Changes.", bold=True)
    _p(doc, (
        "Material changes — including without limitation additional Use Cases, new data "
        "sources, new integrations, or work materially beyond the base scope — require a "
        "written amendment (which may be effected by email exchange between authorised "
        "representatives of the parties) before work on the change commences. The amendment "
        "will describe the change and any resulting adjustment to fees, timeline, or "
        "deliverables."
    ))

    _p(doc, "(c) No Implied Scope.", bold=True)
    _p(doc, (
        "Discussion in working sessions, Weekly Demos, or email correspondence does not "
        "modify the scope of the Services except in accordance with this Section 18. "
        "Neither party is bound by course-of-dealing conduct inconsistent with this "
        "Section 18."
    ))

    # -------- 19. Handoff --------
    _h2(doc, "19. Handoff and Transition")
    _p(doc, (
        "If the parties elect Handoff at the Month 3 Decision Gate, or upon any other "
        "termination where handoff is elected:"
    ))
    _bullet(doc, (
        "Consultant will deliver the Tool, then-current documentation, prompt and "
        "configuration library, and deployment artifacts to Client;"
    ))
    _bullet(doc, (
        "Consultant will train Client's internal champion on Tool operation, including up to "
        "[8] hours of training included within the Pilot Term fees;"
    ))
    _bullet(doc, (
        "Consultant will migrate the Tool's operational configuration to Client's "
        "Subprocessor accounts under Section 11 and will revoke Consultant's operational "
        "access except as permitted under Section 20;"
    ))
    _bullet(doc, (
        "Consultant will provide reasonable transition support, by email and phone, for "
        "fourteen (14) days following handoff at no additional cost. Any additional transition "
        "or operator assistance will be engaged under Section 20."
    ))

    # -------- 20. Continued Operator Services --------
    _h2(doc, "20. Continued Operator Services (Optional)")
    _p(doc, (
        "Prior to handoff under Section 19, Client may elect to engage Consultant for "
        "ongoing operator services in respect of the Tool (“Continued Operator Services”) "
        "under a separate written Statement of Work on terms to be mutually agreed. Continued "
        "Operator Services are not governed by this Agreement, require a separate engagement, "
        "and shall be subject to their own fee schedule, liability terms, and service-level "
        "commitments. Consultant is under no obligation to accept any such engagement. This "
        "Section does not create any continuing obligation on Consultant beyond the Pilot "
        "Term."
    ))

    # -------- 21. Non-Solicit --------
    _h2(doc, "21. Non-Solicitation and Non-Disparagement")
    _p(doc, (
        "(a) During the term of this Agreement and for twelve (12) months thereafter, Client "
        "will not directly or indirectly solicit, hire, or engage any employee, contractor, "
        "or consultant of JIRAH Growth Consulting Ltd. This restriction does not prohibit "
        "general public job postings not directed at Consultant's personnel."
    ))
    _p(doc, (
        "(b) Both parties will refrain from any public or private statements that could "
        "reasonably be construed as defamatory or damaging to the other party's reputation."
    ))

    # -------- 22. Publicity --------
    _h2(doc, "22. Publicity and Case Studies")
    _p(doc, (
        "Consultant may reference the engagement in marketing materials and case studies "
        "using anonymised industry descriptors (for example, “a BC-based building-controls "
        "contractor of comparable size”). Any reference to Client by name, any specific "
        "metric more precise than rounded or banded figures, or any screenshot or "
        "demonstration of the Tool using Client-identifiable data requires Client's prior "
        "written approval, which may be given by email. Client's approval shall not be "
        "unreasonably withheld for anonymised case studies that do not disclose confidential "
        "information."
    ))

    # -------- 23. Dispute Resolution --------
    _h2(doc, "23. Dispute Resolution")
    _p(doc, (
        "The parties will first seek to resolve any dispute through good-faith negotiation. "
        "If unresolved within thirty (30) days, the dispute will proceed to mediation, and "
        "failing resolution by mediation, to binding arbitration conducted in Kelowna, "
        "British Columbia, under the rules of the Arbitration Act (BC). Each party will bear "
        "its own costs unless otherwise determined by the arbitrator."
    ))

    # -------- 24. Governing Law --------
    _h2(doc, "24. Governing Law")
    _p(doc, (
        "This Agreement is governed by and construed in accordance with the laws of the "
        "Province of British Columbia and the federal laws of Canada applicable therein."
    ))

    # -------- 25. Relationship to Strategic Agreement --------
    _h2(doc, "25. Relationship to Strategic Agreement")
    _p(doc, (
        "This Agreement operates independently of the Strategic Execution Advisory Agreement "
        "between the parties dated April 21, 2026. Termination or expiry of one Agreement "
        "does not affect the continued operation of the other. To the extent of any "
        "irreconcilable conflict between this Agreement and the Strategic Agreement on any "
        "subject matter, this Agreement controls with respect to the AI Pilot."
    ))

    # -------- 26. Entire Agreement --------
    _h2(doc, "26. Entire Agreement")
    _p(doc, (
        "This Agreement, together with its Schedules and any written amendments executed in "
        "accordance with Section 18 or Section 27, represents the entire understanding "
        "between the parties regarding the AI Pilot and supersedes all prior discussions, "
        "proposals, or agreements, whether written or oral, on that subject matter."
    ))

    # -------- 27. Amendments --------
    _h2(doc, "27. Amendments")
    _p(doc, (
        "Any modification to this Agreement must be made in writing and signed by authorised "
        "representatives of both parties. An exchange of email between authorised "
        "representatives confirming the terms of a change may constitute a written amendment "
        "for purposes of Section 18 (Change Management)."
    ))

    # -------- Signatures --------
    _p(doc, "")
    _p(doc, "JIRAH Growth Consulting Ltd.")
    _p(doc, "By: ____________________________")
    _p(doc, "Name: Joshua Marshall")
    _p(doc, "Title: Principal & Founder")
    _p(doc, "Date: ___________________________")
    _p(doc, "")
    _p(doc, "Genesis Building Controls Ltd.")
    _p(doc, "By: ____________________________")
    _p(doc, "Name: Trent Novakowski")
    _p(doc, "Title: Owner")
    _p(doc, "Date: ___________________________")

    # -------- Schedule A — IP Matrix --------
    doc.add_page_break()
    _h1(doc, "Schedule A — Intellectual Property Matrix")
    _p(doc, (
        "This Schedule A supplements Section 10 of the Agreement. In the event of any "
        "conflict between Section 10 and this Schedule A, Section 10 controls."
    ))

    _h2(doc, "Use Cases (Section 1, Definitions)")
    _bullet(doc, (
        "Estimating workflow — AI-assisted support for Client's estimating team, covering "
        "[to be scoped jointly during Pilot Week 1] such as quantity takeoff review, pricing "
        "consistency checks, and proposal-narrative assistance."
    ))
    _bullet(doc, (
        "Project-management workflow — AI-assisted support for Client's project-management "
        "team, covering [to be scoped jointly during Pilot Week 1] such as submittal-log "
        "review, RFI triage, and progress-update drafting."
    ))

    _h2(doc, "Ownership and Licence Matrix")

    ip_table = doc.add_table(rows=1, cols=4)
    ip_table.style = "Light Grid Accent 1"
    ip_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = ip_table.rows[0].cells
    for i, text in enumerate(["Asset Category", "Ownership", "Licence Granted", "Examples"]):
        run = header[i].paragraphs[0].add_run(text)
        run.bold = True

    rows = [
        (
            "Reusable Components",
            "Consultant",
            "Perpetual, non-exclusive, royalty-free internal-use licence to Client as "
            "incorporated in the Tool.",
            "Multi-stage validation prompt structures; evaluation-harness methodology; "
            "orchestration templates; abstracted workflow patterns; anonymised lessons "
            "learned.",
        ),
        (
            "Tool (Client-specific configurations, prompts, workflows, integrations)",
            "Client (on payment in full and handoff)",
            "N/A — Client owns outright; Consultant retains the underlying Reusable "
            "Components per row 1.",
            "Genesis-specific estimating prompts referencing Genesis rate cards and "
            "product categories; Genesis PM workflow integrations; Genesis-branded UI; "
            "Genesis-specific evaluation test cases.",
        ),
        (
            "Client Data",
            "Client",
            "Limited licence to Consultant for performance of the Services during the "
            "Pilot Term only.",
            "Estimating history, pricing data, PM records, customer lists, employee "
            "personal information, building-plans and specifications provided by Client.",
        ),
        (
            "AI-Generated Outputs",
            "Unclaimed — neither party asserts copyright",
            "Client may use outputs without restriction as between the parties, subject "
            "to Section 14.",
            "Draft estimates, drafted PM communications, summarisation outputs, extracted "
            "data from Client Data.",
        ),
        (
            "Third-Party Model Weights / Services",
            "Third-party Subprocessor",
            "Pass-through per Subprocessor terms; governed by Section 11.",
            "Anthropic Claude models; OpenAI models (if used); underlying LLM weights and "
            "inference services.",
        ),
        (
            "Handoff Deliverable Package",
            "Client",
            "N/A — delivered to Client at handoff; Section 19.",
            "Source repository, prompt library, configuration files, documentation, "
            "environment-setup instructions, champion training materials.",
        ),
    ]

    for category, ownership, licence, examples in rows:
        row = ip_table.add_row().cells
        row[0].text = category
        row[1].text = ownership
        row[2].text = licence
        row[3].text = examples

    # -------- Schedule B — Subprocessors --------
    doc.add_page_break()
    _h1(doc, "Schedule B — Approved Subprocessors")
    _p(doc, (
        "The Subprocessors listed below are approved for use under the Agreement as of the "
        "Effective Date. Additions are governed by Section 11(b)."
    ))

    sub_table = doc.add_table(rows=1, cols=4)
    sub_table.style = "Light Grid Accent 1"
    sub_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    sub_header = sub_table.rows[0].cells
    for i, text in enumerate(["Provider", "Purpose", "Data Region", "Notes"]):
        run = sub_header[i].paragraphs[0].add_run(text)
        run.bold = True

    subprocessors = [
        (
            "Anthropic, PBC",
            "LLM inference (Claude family)",
            "United States",
            "Commercial API terms apply; no training on input by default under commercial "
            "plans.",
        ),
        (
            "[OpenAI, L.L.C. — if used]",
            "LLM inference and/or embeddings",
            "United States",
            "To be confirmed by Consultant based on final architecture; Client notice "
            "under Section 11(b) if added.",
        ),
        (
            "[Supabase, Inc. — if used]",
            "Database and authentication",
            "Canada (ca-central-1)",
            "To be confirmed by Consultant based on final architecture.",
        ),
        (
            "[Vercel, Inc. — if used]",
            "Hosting and deployment",
            "United States / Canada (region to be configured)",
            "To be confirmed by Consultant based on final architecture.",
        ),
        (
            "[Additional — to be completed by Consultant prior to execution]",
            "",
            "",
            "",
        ),
    ]

    for provider, purpose, region, notes in subprocessors:
        row = sub_table.add_row().cells
        row[0].text = provider
        row[1].text = purpose
        row[2].text = region
        row[3].text = notes

    _p(doc, "")
    _p(doc, (
        "[DRAFTING NOTE FOR COUNSEL: the final Subprocessor list is to be confirmed by "
        "Consultant's technical lead prior to execution based on Plan A / Plan B "
        "architecture selection. Client's written consent under Section 11 extends to the "
        "Subprocessors as finalised in this Schedule B.]"
    ), italic=True)

    # -------- Schedule C — Change Request Process --------
    doc.add_page_break()
    _h1(doc, "Schedule C — Change Request Process")
    _p(doc, "This Schedule C operationalises Section 18 (Change Management).")

    _h2(doc, "Minor Changes")
    _bullet(doc, "Request raised by email from Client's champion to Consultant's lead.")
    _bullet(doc, "Consultant confirms feasibility and effort by reply email.")
    _bullet(doc, "Client confirms by reply email.")
    _bullet(doc, "Change logged in shared change log with date, description, and actor.")
    _bullet(doc, "No fee adjustment; no amendment required.")

    _h2(doc, "Material Changes")
    _bullet(doc, (
        "Request raised by email with description of the change and rationale."
    ))
    _bullet(doc, (
        "Consultant responds within five (5) business days with scope, effort estimate, fee "
        "delta, and timeline impact."
    ))
    _bullet(doc, (
        "Client confirms by reply email from an authorised representative; that exchange "
        "constitutes a written amendment under Section 18(b) and Section 27."
    ))
    _bullet(doc, "Amendment logged and attached to the Agreement as executed.")
    _bullet(doc, "Work on the Material Change does not commence until the exchange is complete.")

    # -------- Save --------
    doc.save(AGREEMENT_PATH)


# ---------------------------------------------------------------------------
# Cover Memo
# ---------------------------------------------------------------------------

def build_cover_memo() -> None:
    doc = new_document()

    _draft_banner(doc)
    _p(doc, "Prepared: April 23, 2026")
    _p(doc, "")

    _h1(doc, "Cover Memo to Counsel — Genesis AI Pilot Agreement")

    _p(doc, (
        "This memo accompanies the attached draft of the Genesis AI Pilot Agreement and "
        "summarises (a) structural decisions already taken by Consultant, (b) items flagged "
        "for counsel's specific review, and (c) items pending external inputs (insurance "
        "broker, Client's registered address, final architecture)."
    ))

    _h2(doc, "1. Structural decisions already taken")
    _bullet(doc, (
        "Standalone agreement, parallel to (not amending) the Strategic Execution Advisory "
        "Agreement dated April 21, 2026. Consultant considered an MSA + SOW structure and "
        "elected standalone for this engagement; see item 8 below for a firm-level question "
        "about whether to adopt MSA + SOW going forward."
    ))
    _bullet(doc, (
        "Three-month pilot at $9,095/month (total $27,285 + GST), with a Day-2 Plan A/B "
        "feasibility gate and a Month 3 decision gate (handoff / extension / termination)."
    ))
    _bullet(doc, (
        "Client owns the final Tool on payment + handoff; Consultant retains Reusable "
        "Components with perpetual internal-use licence to Client. IP Matrix is Schedule A."
    ))
    _bullet(doc, (
        "Client establishes its own third-party API accounts (Anthropic etc.) and bears "
        "usage fees from handoff forward. Consultant uses development accounts during the "
        "Pilot Term and migrates configuration at handoff. Subprocessor list is Schedule B."
    ))
    _bullet(doc, (
        "The “zero-hallucination evaluation harness” is framed as a process commitment, not "
        "an outcome warranty, paired with a human-in-the-loop requirement and an express no-"
        "warranty-of-accuracy clause."
    ))
    _bullet(doc, (
        "Two-tier change management: minor changes by logged email; material changes by "
        "written amendment (email exchange permitted). Process detail is Schedule C."
    ))
    _bullet(doc, (
        "Continued Operator Services are available only under a separate written SOW and "
        "do not extend automatically from the Pilot Agreement."
    ))
    _bullet(doc, (
        "In-contract PIPEDA / PIPA-BC language (no separate DPA schedule) — see item 8 for "
        "a firm-level question about DPA schedule as future template."
    ))

    _h2(doc, "2. Liability cap structure")
    _p(doc, (
        "Section 16(b) proposes a floor + proportional + ceiling structure: greater of "
        "(i) fees paid in 12 months preceding the claim, or (ii) $50,000 CAD, subject to an "
        "overall ceiling of $150,000 CAD. This is intended to (a) preserve meaningful "
        "recovery for Client on small engagements, (b) use the industry-standard 12-month "
        "proportional middle, and (c) hard-cap Consultant's uninsured exposure regardless of "
        "engagement length."
    ))
    _p(doc, (
        "Both the $50,000 floor and $150,000 ceiling are placeholders pending confirmation "
        "of Consultant's Errors & Omissions and cyber-liability coverage limits (see item 4). "
        "The ceiling should be set at or below coverage so uninsured exposure is zero."
    ))

    _h2(doc, "3. Items flagged for counsel's specific attention")
    _bullet(doc, (
        "Section 10 / Schedule A — the definition of “Reusable Components” vs “Tool”. This "
        "is the area most likely to generate a future dispute between the parties. Counsel "
        "is asked to pressure-test the categorical definitions and the worked examples in "
        "Schedule A."
    ))
    _bullet(doc, (
        "Section 14 — confirmation that the evaluation-harness-as-process-commitment plus "
        "human-in-the-loop framing is adequately defensive against a “you promised zero "
        "hallucinations” argument."
    ))
    _bullet(doc, (
        "Section 16(b) liability-cap numbers — dependent on insurance broker input (item 4)."
    ))
    _bullet(doc, (
        "Section 16(f) Subprocessor pass-through — confirmation this construction is "
        "defensible under BC law given Consultant's vendor-selection role."
    ))
    _bullet(doc, (
        "Section 12 Data Protection — confirmation that in-contract construction is "
        "sufficient for Genesis (BC→BC processing), and flag any employee-personal-"
        "information edge cases that may warrant additional treatment."
    ))
    _bullet(doc, (
        "Section 18 + Schedule C change-management two-tier — confirmation that the minor-"
        "vs-material distinction is legally enforceable and does not create an implied-"
        "course-of-dealing problem."
    ))

    _h2(doc, "4. Pending external inputs")
    _bullet(doc, (
        "Insurance broker confirmation — whether Consultant's current E&O policy covers AI-"
        "build / software-development work (many E&O policies exclude it); whether cyber "
        "liability is in place; whether cyber covers Subprocessor-caused breach; current "
        "coverage limits. The Section 16(b) ceiling should sit at or below coverage."
    ))
    _bullet(doc, (
        "Client's registered address — placeholder in the party block; to be confirmed "
        "before execution."
    ))
    _bullet(doc, (
        "Final Subprocessor list (Schedule B) — to be confirmed by Consultant's technical "
        "lead prior to execution based on Plan A / Plan B architecture selection."
    ))
    _bullet(doc, (
        "Client's nominated internal champion — to be named by Client within 14 days of "
        "execution per Section 6; not required pre-signing."
    ))

    _h2(doc, "5. Items intentionally NOT in this draft")
    _bullet(doc, (
        "Uptime, service-level, or maintenance commitments post-handoff — these are excluded "
        "from scope (Section 3) and are available only under a Continued Operator Services "
        "SOW (Section 20)."
    ))
    _bullet(doc, (
        "References to specialist referrals (HR / banking / tax / legal) — these appear in "
        "the parallel Strategic Agreement at the Client's request; not repeated here."
    ))
    _bullet(doc, (
        "Production-grade model-training, fine-tuning, or retraining commitments — excluded "
        "(Section 3)."
    ))
    _bullet(doc, (
        "A separate Data Processing Agreement schedule — decision taken to handle in-"
        "contract for this engagement. See item 8."
    ))

    _h2(doc, "6. Canadian-law specific points for counsel review")
    _bullet(doc, (
        "Copyright of AI-generated outputs — the draft takes the position that neither "
        "party asserts copyright in AI outputs, reflecting the unsettled state of Canadian "
        "law. Counsel is asked to confirm this framing is currently advisable and to flag "
        "any legislative developments that would justify a different approach."
    ))
    _bullet(doc, (
        "PIPEDA / PIPA-BC alignment — Section 12 uses controller/processor language drawn "
        "from PIPEDA with Client as controller and Consultant as service provider. Counsel "
        "is asked to confirm alignment with current OPC and BC OIPC guidance."
    ))
    _bullet(doc, (
        "Arbitration under the Arbitration Act (BC) 2020 — Section 23. Confirm seat and "
        "procedural references are current."
    ))

    _h2(doc, "7. Questions on which counsel's view is requested")
    _bullet(doc, (
        "Is the Section 16(f) Subprocessor pass-through (limited to amounts actually "
        "recovered from Subprocessors, subject to the Section 16(b) caps) defensible, or "
        "should the language be tightened?"
    ))
    _bullet(doc, (
        "Does Section 12 adequately protect Consultant in the event of an employee-data "
        "claim by a Client employee (alleging Client failed to obtain consent)? Is the "
        "Section 16(e)(ii) Client indemnification sufficient?"
    ))
    _bullet(doc, (
        "Is the Section 18 two-tier change-management structure enforceable given the "
        "weekly-demo operational cadence, or does it create an implied course-of-dealing "
        "waiver risk? Should “Material Change” be defined more precisely (e.g., by effort "
        "threshold)?"
    ))
    _bullet(doc, (
        "Is the Section 10(b) assignment (triggered on payment in full and handoff) drafted "
        "correctly to achieve a clean transfer under Canadian copyright law, or should the "
        "assignment mechanics be strengthened (e.g., moral-rights waiver)?"
    ))

    _h2(doc, "8. Firm-level questions (not Genesis-specific)")
    _bullet(doc, (
        "MSA + SOW structure going forward — Consultant anticipates multiple parallel "
        "strategic + AI engagements through 2026. Is it advisable to re-paper the Strategic "
        "Agreement and this AI Pilot Agreement as SOWs under a unified MSA? If yes, at what "
        "point (e.g., at second parallel engagement, at year-end)?"
    ))
    _bullet(doc, (
        "Reusable DPA schedule — should Consultant commission a standalone PIPEDA/PIPA-BC "
        "DPA schedule as a template for future enterprise-style clients who expect a "
        "standalone DPA, even though Genesis does not require one?"
    ))
    _bullet(doc, (
        "Insurance — confirm E&O policy covers AI-build work (common exclusion); confirm "
        "cyber liability is in place; confirm Subprocessor-breach coverage."
    ))

    _p(doc, "")
    _p(doc, "— end of cover memo —", italic=True)

    doc.save(MEMO_PATH)


if __name__ == "__main__":
    os.makedirs(BASE_DIR, exist_ok=True)
    build_agreement()
    build_cover_memo()
    print(f"WROTE: {AGREEMENT_PATH}")
    print(f"       {os.path.getsize(AGREEMENT_PATH)} bytes")
    print(f"WROTE: {MEMO_PATH}")
    print(f"       {os.path.getsize(MEMO_PATH)} bytes")
