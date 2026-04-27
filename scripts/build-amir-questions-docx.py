"""
Build the editable .docx version of the Amir Discovery Call Questions doc.

Mirrors the content of the HTML version, structured as a properly-styled
Word document Joshua can mark up directly. JIRAH brand colours and fonts
mapped to Word-native equivalents (Cambria for display, Calibri for body).
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Brand palette ────────────────────────────────────────────────────────
NAVY = RGBColor(0x0F, 0x26, 0x33)
PLUM = RGBColor(0x38, 0x21, 0x36)
GOLD = RGBColor(0xC7, 0xAD, 0x66)
GOLD_PALE = RGBColor(0xEC, 0xE2, 0xC6)
CREAM = RGBColor(0xF7, 0xF4, 0xE9)
INK = RGBColor(0x0F, 0x26, 0x33)
INK_MID = RGBColor(0x4A, 0x59, 0x68)
INK_LIGHT = RGBColor(0x7B, 0x86, 0x93)
RULE = RGBColor(0xD9, 0xD3, 0xBD)
AMBER = RGBColor(0xB6, 0x7E, 0x2B)

DISPLAY_FONT = "Cambria"
BODY_FONT = "Calibri"

OUT_PATH = Path(
    r"C:\Users\joshu\OneDrive - jirahgrowth.consulting"
    r"\JIRAH Growth Partners - Shared\NEW - JIRAH MASTER"
    r"\01 - Clients\Active\Hatch Interior Design"
    r"\08 - Monthly Retainer"
    r"\2026-04-30 Amir Discovery Call — Questions for Rachel and Shanna.docx"
)


# ─── Helpers ──────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_colour):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


def set_cell_border(cell, side, hex_colour, sz=8):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(sz))
    border.set(qn("w:color"), hex_colour)
    tc_borders.append(border)


def add_run(paragraph, text, *, font=BODY_FONT, size=11, bold=False,
            italic=False, color=INK):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    return run


def heading(doc, text, *, size=20, color=NAVY, space_before=18, space_after=8,
            bottom_border_color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    add_run(p, text, font=DISPLAY_FONT, size=size, bold=True, color=color)
    if bottom_border_color is not None:
        p_pr = p._element.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), bottom_border_color)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
    return p


def body_para(doc, text=None, *, runs=None, size=11, space_after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    if runs:
        for r in runs:
            add_run(p, r["text"], font=r.get("font", BODY_FONT),
                    size=r.get("size", size),
                    bold=r.get("bold", False),
                    italic=r.get("italic", False),
                    color=r.get("color", INK))
    elif text:
        add_run(p, text, size=size)
    return p


def bullet(doc, text=None, *, runs=None, size=11):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    if runs:
        for r in runs:
            add_run(p, r["text"], font=r.get("font", BODY_FONT),
                    size=r.get("size", size),
                    bold=r.get("bold", False),
                    italic=r.get("italic", False),
                    color=r.get("color", INK))
    elif text:
        add_run(p, text, size=size)
    return p


def numbered(doc, text=None, *, runs=None, size=11):
    p = doc.add_paragraph(style="List Number")
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    if runs:
        for r in runs:
            add_run(p, r["text"], font=r.get("font", BODY_FONT),
                    size=r.get("size", size),
                    bold=r.get("bold", False),
                    italic=r.get("italic", False),
                    color=r.get("color", INK))
    elif text:
        add_run(p, text, size=size)
    return p


def callout_box(doc, *, label, body_runs, fill_hex, accent_hex, label_color):
    """Single-cell table styled as a callout."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.0)
    cell = table.cell(0, 0)
    cell.width = Inches(7.0)
    set_cell_bg(cell, fill_hex)
    set_cell_border(cell, "left", accent_hex, sz=24)
    set_cell_border(cell, "top", fill_hex, sz=4)
    set_cell_border(cell, "bottom", fill_hex, sz=4)
    set_cell_border(cell, "right", fill_hex, sz=4)
    # Clear default paragraph and add styled content
    cell.text = ""
    label_p = cell.paragraphs[0]
    add_run(label_p, label, font=BODY_FONT, size=8, bold=True,
            color=label_color)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_before = Pt(4)
    for r in body_runs:
        add_run(body_p, r["text"], font=r.get("font", BODY_FONT),
                size=r.get("size", 12),
                bold=r.get("bold", False),
                italic=r.get("italic", False),
                color=r.get("color", INK))
    # Padding via cell margins (margin in twentieths of a point)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in [("top", 200), ("left", 240),
                      ("bottom", 200), ("right", 240)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)
    # Spacing after the table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return table


def question_block_header(doc, title, tag):
    """Sub-section header for a block of questions."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    add_run(p, title, font=DISPLAY_FONT, size=13, bold=True, color=NAVY)
    add_run(p, "    ")
    add_run(p, tag.upper(), font=BODY_FONT, size=8, bold=True, color=GOLD)
    # Bottom rule
    p_pr = p._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "D9D3BD")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def inline_note(doc, runs):
    """Indented note under a question — used for the commission Q heads-up."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.4)
    pf.space_before = Pt(2)
    pf.space_after = Pt(8)
    p_pr = p._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "B67E2B")
    p_bdr.append(left)
    p_pr.append(p_bdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "ECE2C6")
    p_pr.append(shd)
    for r in runs:
        add_run(p, r["text"], font=BODY_FONT,
                size=r.get("size", 10),
                italic=r.get("italic", True),
                bold=r.get("bold", False),
                color=r.get("color", PLUM))


def header_block(doc):
    """First-page header — JIRAH wordmark left, doc title right."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)

    left = table.cell(0, 0)
    left.width = Inches(3.5)
    left.text = ""
    p = left.paragraphs[0]
    add_run(p, "JIRAH", font=DISPLAY_FONT, size=32, bold=True, color=NAVY)
    sub = left.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    add_run(sub, "GROWTH PARTNERS", font=BODY_FONT, size=8, bold=True,
            color=PLUM)

    right = table.cell(0, 1)
    right.width = Inches(3.5)
    right.text = ""
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(rp, "Discovery Call · Amir", font=DISPLAY_FONT, size=20,
            bold=True, color=NAVY)
    sub = right.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub.paragraph_format.space_before = Pt(0)
    add_run(sub, "HATCH INTERIOR DESIGN", font=BODY_FONT, size=9, bold=True,
            color=PLUM)
    date = right.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date.paragraph_format.space_before = Pt(0)
    add_run(date, "THURSDAY APR 30, 2026 · FOR RACHEL & SHANNA",
            font=BODY_FONT, size=8, color=INK_LIGHT)

    # Bottom rule under the header table
    rule = doc.add_paragraph()
    rule_pr = rule._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0F2633")
    p_bdr.append(bottom)
    rule_pr.append(p_bdr)
    rule.paragraph_format.space_after = Pt(12)


# ─── Build the document ───────────────────────────────────────────────────
def build():
    doc = Document()

    # Page setup — Letter, 0.5" margins
    for section in doc.sections:
        section.page_height = Inches(11.0)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = INK

    # ── Header
    header_block(doc)

    # ── Section 1 — Purpose and Framing
    heading(doc, "Purpose and Framing", size=16, color=NAVY,
            bottom_border_color="C7AD66", space_before=4)

    callout_box(
        doc,
        label="SET THE TONE BEFORE THE CALL",
        fill_hex="0F2633",
        accent_hex="C7AD66",
        label_color=GOLD,
        body_runs=[
            {"text": "This is a discovery call. Nothing more. ",
             "bold": True, "color": GOLD_PALE, "size": 12},
            {"text": "The goal Thursday is to learn what Amir is "
                     "actually proposing, watch how he responds when you "
                     "set the boundaries — and walk away with information, "
                     "not a deal. ", "color": CREAM, "size": 12},
            {"text": "You are not signing anything at this meeting. ",
             "bold": True, "color": GOLD_PALE, "size": 12},
            {"text": "You are not committing to a structure, a timeline, "
                     "or a referral relationship. You're listening, asking, "
                     "and watching.", "color": CREAM, "size": 12},
        ],
    )

    # Posture reminder
    callout_box(
        doc,
        label="QUICK REMINDER ON POSTURE",
        fill_hex="ECE2C6",
        accent_hex="C7AD66",
        label_color=PLUM,
        body_runs=[
            {"text": "Hatch is the established design firm here. ",
             "bold": True, "color": NAVY, "size": 11},
            {"text": "Amir is launching something new and would benefit "
                     "from your reputation to do it. ", "color": INK,
             "size": 11},
            {"text": "He needs you more than you need him. ",
             "italic": True, "color": PLUM, "size": 11},
            {"text": "You don't owe him a decision Thursday — and the best "
                     "partners will give you room to think. Walk in calm, "
                     "curious, and unhurried.", "color": INK, "size": 11},
        ],
    )

    # ── Section 2 — Desired outcomes
    heading(doc, "Desired Outcomes", size=16, color=NAVY,
            bottom_border_color="C7AD66")

    body_para(doc, runs=[
        {"text": "By the end of the call, you want to leave with enough "
                 "information to make a clear-eyed decision in the days "
                 "after. Not a yes. Not a no. ", "size": 11},
        {"text": "A read.", "bold": True, "size": 11, "color": NAVY},
    ], space_after=6)

    bullet(doc, runs=[
        {"text": "Understand exactly what he's proposing. ", "bold": True,
         "color": NAVY},
        {"text": "What does he picture day-to-day? On a typical project, "
                 "what does he expect from you and what does he commit to "
                 "do himself?"},
    ])
    bullet(doc, runs=[
        {"text": "Read his response when you name the ethical boundary. ",
         "bold": True, "color": NAVY},
        {"text": "No commissions, no kickbacks — that's a RID code-of-ethics "
                 "issue. How he reacts to that bright line tells you who "
                 "he is in harder moments."},
    ])
    bullet(doc, runs=[
        {"text": "Verify the offer is real. ", "bold": True, "color": NAVY},
        {"text": "Named projects on the East Coast. Named dentists. Named "
                 "contacts at Henry Schein, Patterson, Sinclair. If it's "
                 "vague, it's a sales pitch — not a pipeline."},
    ])
    bullet(doc, runs=[
        {"text": "Decide whether a small trial project is worth exploring. ",
         "bold": True, "color": NAVY},
        {"text": "Not a partnership. Not an MOU. One real project where "
                 "you both show up, do the work, and see how it goes."},
    ])
    bullet(doc, runs=[
        {"text": "Decide whether a follow-up conversation makes sense in "
                 "2–3 weeks. ", "bold": True, "color": NAVY},
        {"text": "Or whether the answer is a polite \"thanks, this isn't "
                 "going to fit.\""},
    ])

    body_para(doc, runs=[
        {"text": "What is "},
        {"text": "not", "italic": True},
        {"text": " a desired outcome: ", "bold": True, "color": NAVY},
        {"text": "a signed agreement, a verbal commitment to a structure, "
                 "an exclusivity carve-out for any region or vertical, or "
                 "shared access to your client list, processes, or channel "
                 "relationships. None of that should leave the room "
                 "Thursday."},
    ], space_after=10)

    # ── Page break before questions
    doc.add_page_break()

    # ── Section 3 — Questions
    heading(doc, "Questions to Bring Into the Room", size=16, color=NAVY,
            bottom_border_color="C7AD66", space_before=0)

    body_para(doc, runs=[
        {"text": "Pick what fits the flow of the conversation. You don't "
                 "need to fire all of these — but having them in the back "
                 "pocket keeps the meeting on track. Listen carefully to "},
        {"text": "how", "italic": True},
        {"text": " he answers, not just what he says."},
    ], space_after=8)

    # Block 1 — Open it up
    question_block_header(doc, "Open it up", "What is he actually proposing")
    numbered(doc, "Walk us through what you have in mind. What does a "
                  "typical project look like under this partnership?")
    numbered(doc, "Tell us about your dental design company — what's it "
                  "called, what's the positioning, who's doing the design "
                  "work?")

    # Block 2 — Verify the offer
    question_block_header(doc, "Verify the offer",
                          "Is this real or aspirational")
    numbered(doc, "Your East Coast work — how did that come about? Is it "
                  "one project closed, or is there a real pipeline behind "
                  "it?")
    numbered(doc, "What are your relationships with Henry Schein, Patterson, "
                  "and Sinclair? Where are they strongest?")

    # Block 3 — Test the structure
    question_block_header(doc, "Test the structure",
                          "How does this actually work")
    numbered(doc, "How are you thinking about the financial relationship? "
                  "Referral fees, percentage-of-project, each firm bills "
                  "its own client?")
    inline_note(doc, [
        {"text": "Heads up: ", "italic": False, "bold": True, "color": PLUM,
         "size": 10},
        {"text": "we can't accept commissions or referral fees — RID code "
                 "of ethics. Watch how he reacts to that boundary. ",
         "italic": True, "color": PLUM, "size": 10},
        {"text": "Green:", "italic": False, "bold": True, "color": PLUM,
         "size": 10},
        {"text": " \"Understood, let's figure out a structure that works.\" ",
         "italic": True, "color": PLUM, "size": 10},
        {"text": "Yellow:", "italic": False, "bold": True, "color": PLUM,
         "size": 10},
        {"text": " \"Hmm, let me think about that.\" ", "italic": True,
         "color": PLUM, "size": 10},
        {"text": "Red:", "italic": False, "bold": True, "color": PLUM,
         "size": 10},
        {"text": " \"Come on, everyone does it\" or any pivot to relabeling "
                 "the money.", "italic": True, "color": PLUM, "size": 10},
    ])
    numbered(doc, "Who owns the client relationship? If a dentist comes "
                  "through one of us, who's on the contract going forward?")
    numbered(doc, "How would each of us be represented in marketing? "
                  "Co-branded, or each firm with its own role?")

    # Block 4 — Test the fit
    question_block_header(doc, "Test the fit",
                          "Does this complement or complicate")
    numbered(doc, "Are you imagining any kind of exclusivity — vertical, "
                  "geographic, or time-bound?")
    numbered(doc, "How would you feel about Hatch continuing to deepen our "
                  "own relationships with Sinclair, Henry Schein, and other "
                  "GCs? Inside this arrangement or outside it?")

    # Block 5 — Close cleanly
    question_block_header(doc, "Close cleanly", "Leave room to think")
    numbered(doc, "What's your timeline on this? When are you hoping to "
                  "know one way or the other?")
    numbered(doc, "If there's a path forward, what would you want to see "
                  "at a follow-up conversation in 2–3 weeks?")

    # ── Closing note
    callout_box(
        doc,
        label="BEFORE YOU WALK OUT",
        fill_hex="F7F4E9",
        accent_hex="382136",
        label_color=PLUM,
        body_runs=[
            {"text": "Don't commit to anything Thursday. ", "bold": True,
             "color": NAVY, "size": 11},
            {"text": "Not a structure, not a fee model, not exclusivity, "
                     "not a timeline. The cleanest exit line: ", "size": 11},
            {"text": "\"This is helpful — we'd like to take it away, talk "
                     "it through, and come back to you with our read.\" ",
             "italic": True, "color": PLUM, "size": 11},
            {"text": "That gives you the time to debrief, run the answers "
                     "through the rubric Joshua walked you through, and "
                     "decide what (if anything) is worth a next step. The "
                     "right partner will respect that. The wrong one will "
                     "push.", "size": 11},
        ],
    )

    # Footer note
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(18)
    add_run(footer_p, "JIRAH · AMIR DISCOVERY QUESTIONS · HATCH · "
                      "PREPARED BY JOSHUA MARSHALL · jirahgrowth.com",
            font=BODY_FONT, size=7, bold=True, color=INK_LIGHT)

    # Save
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
