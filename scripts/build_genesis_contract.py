"""One-off: generate the Genesis Strategic Execution Advisory Agreement (.docx).

Mirrors the Hatch Monthly Coaching Agreement skeleton, adapted for:
- 12-month fixed initial term
- $15,155/mo + GST
- Strategic execution advisory (AI carved out, to be covered under a separate agreement)
- Explicit out-of-scope clauses for HR, banking, legal, tax
"""

from __future__ import annotations

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt

sys.path.insert(0, str(Path(__file__).parent))
from _paths import ACTIVE_CLIENTS, assert_no_legacy_segment

OUT_PATH = str(
    ACTIVE_CLIENTS
    / "Genesis Systems"
    / "03 - Admin"
    / "Contract & DocuSign"
    / "Genesis_JIRAH Strategic Execution Advisory Agreement.docx"
)
assert_no_legacy_segment(OUT_PATH)


def build() -> str:
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def h1(text: str) -> None:
        doc.add_paragraph(text, style="Heading 1")

    def h2(text: str) -> None:
        doc.add_paragraph(text, style="Heading 2")

    def para(text: str = "") -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    h1("Strategic Execution Advisory Agreement")

    para("Effective Date: April 21, 2026")
    para("")

    para("JIRAH Growth Consulting Ltd.")
    para("3610 Weston Rd, Kelowna, BC V1X 6K2 (“Consultant”)")
    para("")
    para("And")
    para("")
    para("Genesis Building Controls Ltd.")
    para("[Registered address — to be confirmed] (“Client”)")

    h2("1. Engagement")
    para(
        "The Client retains the Consultant to provide strategic execution advisory services across a "
        "12-month engagement horizon, supporting the execution of the Client’s multi-region "
        "growth plan. The Consultant will act as an extended executive partner working alongside the "
        "Client’s ownership and leadership team to advance the priorities identified in the "
        "Strategic Planning Report and Alberta Expansion Preparation Plan previously delivered to the "
        "Client."
    )
    para(
        "The parties acknowledge this is a best-efforts, outcomes-oriented engagement. The Consultant "
        "will allocate additional time beyond the planned monthly allocation at the Consultant’s "
        "discretion where required to advance committed milestones, without additional fees."
    )

    h2("2. Scope of Services")
    para("The Consultant will provide:")
    bullet(
        "Weekly working sessions with the Client’s ownership and operating team (virtual or "
        "in-person) to drive execution against the 12-month roadmap;"
    )
    bullet(
        "Strategic execution support across the six (6) Action Plans identified in the Strategic "
        "Planning Report, including operations infrastructure, regional-manager model, estimating "
        "capacity, and Alberta market-entry readiness;"
    )
    bullet(
        "Weekly progress tracking and action-register maintenance, including Last-Week / This-Week / "
        "Next-Two-Week visibility;"
    )
    bullet("Monthly written check-in against 12-month milestones (Month 3 / 6 / 9 / 12);")
    bullet("Quarterly step-back reviews to re-align direction against the plan;")
    bullet(
        "Facilitation of strategic conversations with internal stakeholders (e.g., Vancouver operations, "
        "prospective Alberta hire) where scoped into the weekly cadence;"
    )
    bullet("Reporting, process management, and accountability support related to agreed priorities;")
    bullet("Email and/or phone availability between sessions as reasonably required.")

    para(
        "The following are expressly out-of-scope under this Agreement and, where required, will be "
        "covered under a separate agreement or referred to specialist resources identified by the "
        "Client:"
    )
    bullet(
        "AI integration and pilot development (to be covered under a separate written agreement between "
        "the parties);"
    )
    bullet("HR policy drafting, employee-handbook buildout, and employment-law advisory;")
    bullet("Banking, lending, tax, and accounting advisory;")
    bullet("Legal drafting or counsel.")

    para(
        "Any additional work beyond the agreed scope will be subject to mutual written consent and may "
        "be engaged under a separate Statement of Work or written amendment."
    )

    h2("3. Term")
    para(
        "This Agreement shall commence on April 21, 2026 (the “Effective Date”) and continue "
        "for an initial term of twelve (12) months. Upon expiry of the initial term, the Agreement will "
        "continue on a month-to-month basis unless either party provides written notice of non-renewal "
        "or the parties execute a renewal agreement. Either party may terminate this Agreement in "
        "accordance with Section 9."
    )

    h2("4. Fees and Payment")
    bullet("Monthly fee: $15,155 CAD + applicable GST.")
    bullet(
        "Invoiced monthly, in advance, with an accompanying hours-detail summary for the prior month."
    )
    bullet("Payment due upon receipt of invoice.")
    bullet(
        "Late payments exceeding 30 days may result in suspension of services and/or interest charged "
        "at 2% per month (24% per annum)."
    )
    bullet("No onboarding or setup fees apply.")
    bullet(
        "Fees are fixed for the duration of the initial 12-month term. Upon renewal, fees are subject "
        "to review with 30 days’ written notice."
    )

    h2("5. Meeting Expectations")
    para(
        "Meetings shall be conducted on a weekly cadence (virtual or in-person, as agreed "
        "session-by-session), with a monthly written check-in and quarterly in-person step-back. The "
        "Consultant may include additional members of the JIRAH team (e.g., co-founder or "
        "subject-matter contributors) in relevant sessions at no additional cost to the Client. The "
        "Consultant will remain available for reasonable communication via email and/or phone between "
        "sessions."
    )

    h2("6. Confidentiality")
    para(
        "Both parties agree to maintain the confidentiality of all proprietary, personal, or business "
        "information disclosed during the course of this Agreement. Each party shall take reasonable "
        "precautions to protect such information and shall not disclose it to any third party without "
        "prior written consent, except as required by law. This clause survives termination of the "
        "Agreement."
    )

    h2("7. Intellectual Property")
    para(
        "All materials, templates, frameworks, and proprietary methodologies provided by the Consultant "
        "remain the sole property of the Consultant. The Client receives a non-exclusive, "
        "non-transferable right to use such materials internally for its own business purposes. All "
        "Client materials, data, and information remain the property of the Client."
    )

    h2("8. Liability and Indemnification")
    para(
        "The Consultant agrees to perform services with professional skill and care and maintains "
        "Errors & Omissions insurance. The Consultant shall not be liable for any indirect, incidental, "
        "consequential, or special damages, including loss of revenue, profit, or goodwill, arising "
        "from the services provided under this Agreement. The Client agrees to indemnify and hold "
        "harmless the Consultant from any claims or liabilities resulting from reliance on advice or "
        "deliverables provided, except in cases of gross negligence or willful misconduct."
    )

    h2("9. Termination")
    para(
        "Either party may terminate this Agreement by providing 30 days’ written notice. In the "
        "event of a material breach that remains uncured 15 days following written notice of the "
        "breach, the non-breaching party may terminate immediately. All fees and expenses incurred up "
        "to the effective date of termination shall remain payable. Sections 6, 7, 8, 10, and 11 shall "
        "survive termination."
    )

    h2("10. Non-Solicitation & Non-Disparagement")
    para(
        "During the term of this Agreement and for 12 months thereafter, the Client agrees not to "
        "directly or indirectly solicit, hire, or engage any employee, contractor, or consultant of "
        "JIRAH Growth Consulting Ltd. Both parties further agree to refrain from any public or private "
        "statements that could reasonably be construed as defamatory or damaging to the other’s "
        "reputation."
    )

    h2("11. Publicity & Case Studies")
    para(
        "The Client grants the Consultant permission to reference the engagement in marketing materials "
        "or case studies, provided no confidential or proprietary information is disclosed without "
        "written consent."
    )

    h2("12. Dispute Resolution")
    para(
        "The parties shall first seek to resolve any dispute through good-faith negotiation. If "
        "unresolved, the dispute shall proceed to mediation, and failing that, to binding arbitration "
        "conducted in Kelowna, British Columbia, under the rules of the Arbitration Act (BC). Each "
        "party shall bear its own costs unless otherwise determined by the arbitrator."
    )

    h2("13. Governing Law")
    para(
        "This Agreement shall be governed by and construed in accordance with the laws of the Province "
        "of British Columbia and the federal laws of Canada applicable therein."
    )

    h2("14. Entire Agreement")
    para(
        "This Agreement, together with the Strategic Planning Report and Alberta Expansion Preparation "
        "Plan previously delivered to the Client (which together describe the execution roadmap), "
        "represents the entire understanding between the parties with respect to the subject matter "
        "hereof and supersedes all prior discussions, proposals, or agreements, whether written or oral."
    )

    h2("15. Amendments")
    para("Any modification to this Agreement must be made in writing and signed by both parties.")

    para("")
    para("JIRAH Growth Consulting Ltd.")
    para("By: ____________________________")
    para("Name: Joshua Marshall")
    para("Title: Principal & Founder")
    para("Date: ___________________________")
    para("")
    para("Genesis Building Controls Ltd.")
    para("By: ____________________________")
    para("Name: Trent Novakowski")
    para("Title: Owner")
    para("Date: ___________________________")

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(f"WROTE: {path}")
    print(f"SIZE:  {os.path.getsize(path)} bytes")
